from docx import Document
from openpyxl import Workbook
from app.models.resume import Resume

def export_to_word(resume: Resume, filename: str):
    """Экспорт резюме в Word"""
    doc = Document()
    doc.add_heading(f"{resume.secondname} {resume.firstname} {resume.patronymic or ''}", 0)
    doc.add_paragraph(f"Email: {resume.email}")
    doc.add_paragraph(f"Телефон: {resume.phone}")
    doc.add_paragraph(f"Должность: {resume.position}")
    doc.save(filename)
    return filename

def export_to_excel(resumes: list, filename: str):
    """Экспорт списка резюме в Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Резюме"
    
    # Заголовки
    ws.append(["ID", "Паспорт", "ФИО", "Должность", "Email", "Телефон", "Статус"])
    
    # Данные
    for resume in resumes:
        fio = f"{resume.secondname} {resume.firstname} {resume.patronymic or ''}"
        ws.append([
            resume.id,
            resume.passport,
            fio,
            resume.position,
            resume.email,
            resume.phone,
            "Активно" if resume.status else "Не активно"
        ])
    
    wb.save(filename)
    return filename
