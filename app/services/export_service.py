from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
# УДАЛИТЬ ЭТИ СТРОКИ:
# from reportlab.lib.pagesizes import A4
# from reportlab.pdfgen import canvas
from typing import List
import os
from datetime import datetime
from app.models.resume import Resume

class ExportService:
    
    @staticmethod
    def export_resume_to_word(resume: Resume, filename: str = None) -> str:
        if not filename:
            filename = f"exports/word/resume_{resume.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        os.makedirs(os.path.dirname(filename), exist_ok=True)
        
        doc = Document()
        doc.add_heading(f"{resume.secondname} {resume.firstname} {resume.patronymic or ''}", 0)
        doc.add_paragraph(f"Email: {resume.email}")
        doc.add_paragraph(f"Телефон: {resume.phone}")
        doc.add_paragraph(f"Должность: {resume.position}")
        doc.save(filename)
        return filename
    
    @staticmethod
    def export_resumes_to_excel(resumes: List[Resume], filename: str = None) -> str:
        if not filename:
            filename = f"exports/excel/resumes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        os.makedirs(os.path.dirname(filename), exist_ok=True)
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Резюме"
        
        headers = ["ID", "Паспорт", "ФИО", "Должность", "Email", "Телефон", "Статус"]
        ws.append(headers)
        
        for cell in ws[1]:
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='center')
        
        for resume in resumes:
            fio = f"{resume.secondname} {resume.firstname} {resume.patronymic or ''}"
            ws.append([
                resume.id,
                resume.passport,
                fio,
                resume.position,
                resume.email,
                resume.phone,
                "Активно" if resume.status else "Не активно"
            ])
        
        wb.save(filename)
        return filename
